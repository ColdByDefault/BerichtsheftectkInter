#ColdByDefault

# importrs
import tkinter as tk
import customtkinter as ctk
from tkinter import messagebox
from docx import Document
from datetime import datetime


#Notes_dict
notes = {
    # LF-00 Einführung
    'LF-00': {
        'montag': """Cold Auftaktveranstaltung, Kennenlernrunde, Organisatorisches
Ausblick Inhalte LF00, Grundlagen Teams, Der Support-Bereich im Lernsystem,
Das RemoteLab Arbeitstechniken und Methoden, Die Pomodoro Technik,
Das Pareto-Prinzip, Das Eisenhower-Prinzip, Ausblick auf Tag 2,
Methoden: Präsentation, Diskussion, Videos, F&A""",
        'dienstag': """Frage- und Antwortrunde, Informationen beschaffen und
verwerten, Quellen, Suchmaschinen, Informationen zur Ausbildung, Prüfungsfragen,
Prüfungen, Windows und Kommando-Zeile, Windows Tipps, Tastaturkürzel,
Die Kommandozeile, Suchen in Windows, Dateiendungen, Review des Tages,
Ausblick auf Tag 3, Methoden: Präsentation, Einzelarbeit, Gruppenarbeit,
Diskussion, Präsentation der TN, F&A""",
        'mittwoch': """Wiederholung des Vortags, Frage- und Antwortrunde,
Fachbuch und Fachartikel bearbeiten, Der BASIS-Plan, Einen Fachartikel zum
Thema Virtualisierung bearbeiten Mindmap erstellen und nutzen, Was ist
ein Mindmap? Beispiele, unterschiedliche online-tools Aufgabe zur Mindmap,
Methoden: Präsentation, Einzel-/Gruppenarbeit, Präsentation der Ergebnisse
durch TN, Diskussion, F&A""",
        'donnerstag': """Review, Grundlagen Mathematik für Informatiker:
Dreisatz (proportional, anti-), Gleichungen (Klammern, Variable),
Zahlensysteme (Dezimal, Binär, Hexadezimal), Aufgaben zu den drei Themen, Q&A;
Einfache Präsentation: Aufbau, Todsünden; Ausblick auf Tag 5 mit Methoden:
Präsentation, Einzel-/Gruppenarbeit, Ergebnispräsentation durch die TN,
Fragerunden, Videos.""",
        'freitag': """Wiederholung, F&A, Abschlussaufgabe: Vorbereitung,
Präsentation der Teilnehmer:innen, Feedback. Methoden: Gruppenarbeit,
Präsentationen, Feedback."""
    },
    # LF-01 Unternehmen & Eigne Rolle
    'LF-01': {
        'montag': """Allgemeine Einführung für alle Kursteilnehmer, Agenda Lernfeld
„LF01V2“, PowerPoint 01 Kapitel 1.1 Lernziele LF 01 Duales Ausbildungssystem
Fortbildung/Weiterbildung Übungen: Kompetenzcheck: Seite 16 AB: Seite 12;
Aufgabe 1.1 Kompetenzcheck: Seite 18; 1 PowerPoint 02 - Kapitel 1.2 Rechte und
Pflichten im Arbeitsleben Review des Tages, Ausblick auf den nächsten Tag,
Unterricht mit Folien, Beispielen und Internetseiten geführt, Beispiele
angewendet, Einzel- und Gruppenarbeiten, Lehrgespräch geführt.""",
        'dienstag': """+ Wiederholung des Unterrichtsstoffs vom Vortag Frage- und
Antwortrunde Sozialversicherungen Steuerklassen Ausbildungsvergütung,
Teilzeitausbildung Übungen Kompetenzcheck: Seite 27; Aufgabe 1 AB: Seite 32;
Aufgabe 11.1 + 2, Aufgabe 12.1 - 5 Arbeitsrechte, Kündigung Übungen AB: Seite 25;
Aufgabe 8.1 Mitbestimmung der Arbeitnehmer/innen Betriebsrat/JAV
Betriebsvereinbarung Tarifvertrag, Inhalte, Verhandlungen, Schlichtung,
Arbeitskampf Review des Tages Ausblick auf den nächsten Tag Unterricht mit
Folien, Beispielen und Internetseiten geführt, Beispiele angewendet, Einzel- und
Gruppenarbeiten, Lehrgespräch geführt.""",
        'mittwoch': """Wiederholung des Unterrichtsstoffs vom Vortag PowerPoint 03
Kapitel 1.3 Arbeitsteilung, Globalisierung Unternehmen, Betrieb Verflechtungen
Produktionsfaktoren, -sektoren Unternehmensziele, Compliance PowerPoint 04
Kapitel 1.3.4 Handlungsvollmachten und Prokura Review des Tages, Ausblick auf den
nächsten Tag Unterricht mit Folien, Beispielen und Internetseiten gemacht,
Beispiele angewendet, Einzel- und Gruppenarbeiten durchgeführt,
Lehrgespräch geführt.""",
        'donnerstag': """Aufbauorganisation, (Video: Aufbauorganisation) Aufgabenanalyse,
Stellenbildung Führungsstile Leitungssysteme PowerPoint 05 Kapitel 1.3.6
Produktionsfaktoren Güterarten Ökonomische Prinzipien Wirtschaftskreislauf
Marktformen E-Business, E-Commerce Vollkommener Markt Monopol, etc Kompetenzcheck
Seite 98 1 + 2 Preisbildung Käufer-, Verkäufermarkt Kundenstruktur, -verhalten
Review des Tages, Ausblick auf den nächsten Tag Unterricht mit Folien, Beispielen und
Internetseiten gemacht, Beispiele angewendet, Einzel- und Gruppenarbeiten durchgeführt,
Lehrgespräch geführt.""",
        'freitag': """+ Wiederholung des Unterrichtsstoffs der Woche Frage- und Antwortrunde
ABC-Kundenanalyse Konjunkturphasen LZK Review der Woche Unterricht mit Folien,
Beispielen und Internetseiten gemacht, Beispiele angewendet, Einzel- und
Gruppenarbeiten durchgeführt, Lehrgespräch geführt."""
    },
    # LF-02 Arbeitzplatz nach Kundenwunsch
    'LF-02': {
        'montag': """Begrüßung, Einführung in das LF. Zusammenarbeit im Kurs vorbereiten.
2.1.1 Eine Einführung in Grundfunktionen des Computers geben. 2.1.2 Bedeutende
Entwicklungsschritte in der Computertechnik. 2.1.3 Entwicklungstrends präsentieren.
Exkurs: Diagramme in PowerPoint erstellen. 2.1.4 Komponentenhersteller und
Systemarchitekturen präsentieren. Aufgaben zur Übung.""",
        'dienstag': """- 2.2.1 Arbeitsplätze und Arbeitsumgebungen für IT-Systeme beschreiben
- 2.2.2 Marktgängige IT-Systeme vorstellen - 2.2.3 Das Leistungsportfolio im IT-Bereich
präsentieren - 2.3.1 Qualität und Leistungsfähigkeit von IT-Systemen und IT-Services
beschreiben - Aufgaben zum Üben""",
        'mittwoch': """- 2.3.2 Umweltschutz und Green-IT als wichtige IT-Ziele darstellen
- 2.3.3 Wirtschaftlichkeit von IT-Systemen erläutern - 2.3.4 IT-Sicherheit von IT-Systemen,
Informations- und Datenschutz erläutern - 2.4.1 Zentraleinheit, Mainboard und
Betriebssystem unterscheiden - Aufgaben zum Üben""",
        'donnerstag': """- 2.4.2 Hauptplatine, Mainboard und die Komponenten beschreiben
- 2.4.3 Prozessoren genauer beschreiben - 2.4.4 Arbeitsspeicher, RAM-Speicher erläutern und
unterscheiden - Aufgaben zum Üben""",
        'freitag': """ 2.4.4 Arbeitsspeicher, RAM-Speicher erläutern und unterscheiden
- 2.4.5 Schnittstellen und Anschlüsse am Mainboard erläutern - 2.4.6 Netzteile beschreiben und
unterscheiden - 2.4.7 Festplattenarten unterscheiden und erläutern - 2.4.8 Tastaturen
unterscheiden und präsentieren - 2.4.9 Monitore vergleichen und präsentieren
- Aufgaben zum Üben"""
    },
    #'LF-03':{'montag':""" """,'dienstag':""" """,'mittwoch':""" """,'donnerstag':""" """,'freitag':""" """},
    #'LF-04':{'montag':""" """,'dienstag':""" """,'mittwoch':""" """,'donnerstag':""" """,'freitag':""" """},
    #'LF-05':{'montag':""" """,'dienstag':""" """,'mittwoch':""" """,'donnerstag':""" """,'freitag':""" """},
    #'LF-06':{'montag':""" """,'dienstag':""" """,'mittwoch':""" """,'donnerstag':""" """,'freitag':""" """},
    #'LF-07':{'montag':""" """,'dienstag':""" """,'mittwoch':""" """,'donnerstag':""" """,'freitag':""" """},
    #'LF-08':{'montag':""" """,'dienstag':""" """,'mittwoch':""" """,'donnerstag':""" """,'freitag':""" """},
    #'LF-09':{'montag':""" """,'dienstag':""" """,'mittwoch':""" """,'donnerstag':""" """,'freitag':""" """},
    #'LF-10':{'montag':""" """,'dienstag':""" """,'mittwoch':""" """,'donnerstag':""" """,'freitag':""" """},
    'LF-F1': {
        'montag': """Online Practice: English 4 IT Wiederholung Passive Describing a Process
IT Milestones Chapter 4 (Bios and operating systems) pages Vorbereitung in
Gruppenarbeit Austausch und Korrektur in der großen Gruppe Speaking Exercise
(IT Quiz) Paar-Arbeit IT Milestones Chapter 6""",
        'dienstag': """ finish IT Milestones Chapter 6 Networks pages input If- clauses
If-clause Übungen If-clause Lösungen besprechen Websites um if-clauses zu üben
IT Milestones Chapter 9 Writing software pages Kahoot for single players""",
        'mittwoch': """ Presentation Themes IT Milestones Module G - Presentations
gemeinsam Lösungen überprüfen Evaluationskriterien Abschlusspräsentation besprochen
IT-Milestones Chapter 11 Enterprise Software Präsentationen vorbereiten und besprechen
Quizlet - Vocabulary Chapters 1 - 12 """,
        'donnerstag': """ finish IT Milestones Chapter 11 correct answers together
IT Milestones Module D Written Communication Written Correspondence Writing an enquiry,
an offer, and an order Abschlusspräsentationen weiter fertigstellen Kahoot (computer science) """,
        'freitag': """ IT-Milestones Online Grammar training Modules D - F Test 2 Korrektur
besprochen Modules A - H in Gruppenarbeit Modules.pdf Kurzpräsentationen der Module
Fertigstellung der schriftlichen Präsentationen Abgabe der schriftlichen Präsentationen
Quizlet Chapters 1 - 12 """},
    'LF-F2': {
        'montag': """ Begrüßung und Organisatorisches Vorschau auf die Kommende Woche
Wünsche und Ziele Bewerbungsunterlagen: Lebenslauf und Anschreiben Übungen Zeit zum
Erstellen der Unterlagen """,
        'dienstag': """ Selbstmarketing und USP Übung und Diskussion zu Selbstreflektion
Wertequadrat Gruppenübung und Diskussion zu Selbstbewusstsein Internetmarketing USP Zeit
zum Erstellen der Bewerbungsunterlagen """,
        'mittwoch': """ Online Bewerbungen Unterschied Digital/Papier-Bewerbungen Fragebögen
Bewerbungen per E-Mail Praktikumsplatzsuche Zeit zum Erstellen der Bewerbungsunterlagen und
Prakitkumsplatzsuche """,
        'donnerstag': """ Bewerbungsgespräche Vorbereitung aufs Gespräch Präsentation und
Selbstdarstellung "Stärken stärken, Schwächen schwächen Zeit zum Erstellen der
Bewerbungsunterlagen und Praktikumsplatzsuche """,
        'freitag': """ Vertiefung Online Werkzeuge LinkedIn Einführung Tipps zu Profilerstellung
und Jobsuche Transparenz auf dem Arbeitsmarkt: Kununu und co Gehaltsverhandlungen Feedback
und Abschluss Kahoot Quiz als Wiederholung/Vertiefung| """},
    'LF-F3': {
        'montag': """ Vorstellung und Einführung Grundlagen Microsoft Office mit Übungen
(Grundfunktionen, Erstellen eines Posters) Grundlagen Microsoft PowerPoint mit Übungen
(Grundfunktionen, Dos and Donts) Feedback Selbstlernphase zur Vertiefung """,
        'dienstag': """ Grundlagen Microsoft Excel Grundfunktionen Formatierung Übungen zur
SUMMEN-Funktion WENN-Funktion Verschachtelte Funktion Spiele in Excel Selbstlernphase zum
Abschluss """,
        'mittwoch': """ Projekt: "Hacking-Office" Selbstständige Bearbeitung von Projektvorhaben
in Microsoft Office In zwei verschiedenen Programmen soll je ein Vorhaben umgesetzt werden.
Dieses kann klassisch sein, also standardmäßige Nutzung der Software oder "experimentell",
wo die Software kreativ genutzt wird um Ziele zu erreichen. In regelmäßigen Abständen wurde
der Fortschritt überprüft und ggf. Hilfestellung geleistet. """,
        'donnerstag': """ Selbstständige Bearbeitung von Projektvorhaben in Microsoft Office
In zwei verschiedenen Programmen soll je ein Vorhaben umgesetzt werden. Dieses kann
klassisch sein, also standardmäßige Nutzung der Software oder "experimentell", wo die
Software kreativ genutzt wird um Ziele zu erreichen. In regelmäßigen Abständen wurde der
Fortschritt überprüft und ggf. Hilfestellung geleistet Präsentation der Projekte.| """,
        'freitag': """ Präsentationen der Projekte Diskussion und Fragerunde Feedback und
Abschluss Kahoot Quiz als Wiederholung Vertiefung und Wiederholung durch Selbstlernphase """},

}




# Create the main window using CustomTkinter's CTk class
root = ctk.CTk()
root.title("Berichtsheft GFN/HD 2024")

# Apply a CustomTkinter theme
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("dark-blue")

root.geometry("600x800")  # Adjust window size as needed
root.resizable(False, False)  # Window non-resizable

# Function to switch frames
def raise_frame(frame):
    frame.tkraise()

# Create frames for each tab content using CustomTkinter CTkFrame
frame1 = ctk.CTkFrame(root)
frame2 = ctk.CTkFrame(root)

# Position the frames in the same grid location; they will be stacked
for frame in (frame1, frame2):
    frame.grid(row=1, column=0, columnspan=4, sticky='nsew', padx=20, pady=20)

root.columnconfigure(0, weight=1)
root.columnconfigure(1, weight=1)
root.columnconfigure(2, weight=1)
root.columnconfigure(3, weight=1)
root.rowconfigure(1, weight=1)

# Tab buttons to switch frames
tab1_button = ctk.CTkButton(root, text='Allgemeine Information', command=lambda: raise_frame(frame1), width=200, corner_radius=10, hover=True)
tab2_button = ctk.CTkButton(root, text='Tägliche Notes', command=lambda: raise_frame(frame2), width=200, corner_radius=10, hover=True)

tab1_button.grid(row=0, column=0, pady=8, padx=50)
tab2_button.grid(row=0, column=2, pady=8, padx=50)

def update_text_fields():
    selected_lf = replacements['[lf_num]'].get()
    schedule = notes.get(selected_lf, {})
    for i, day in enumerate(['montag', 'dienstag', 'mittwoch', 'donnerstag', 'freitag']):
        text_boxes[i].delete("1.0", "end")
        text_boxes[i].insert("1.0", schedule.get(day, "Keine Information"))

# Function to find and replace text in the Word document
def find_and_replace(document, replacements):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        placeholder = f'[{key.strip("[]")}]'
                        if placeholder in paragraph.text:
                            print(f"Replacing '{placeholder}' with '{value.get()}'")
                            paragraph.text = paragraph.text.replace(placeholder, value.get())

# Function to collect text from tk.Text and update replacements
def update_replacements_with_text():
    daily_vars = ['[text1]', '[text2]', '[text3]', '[text4]', '[text5]']
    for var, text_box in zip(daily_vars, text_boxes):
        text_content = text_box.get("1.0", "end-1c")  # Get text content
        replacements[var].set(text_content)  # Update the corresponding StringVar
        
# Function to save the updated document
def save_updated_document():
    try:
        update_replacements_with_text()  # Update replacements dictionary 
        
        selected_lf_num = replacements['[lf_num]'].get()
        current_date = datetime.now().strftime("%d%m%Y")
        file_name = f"{current_date}{selected_lf_num}.docx"
        doc = Document('pythonCode\\Berichtshefte.docx')
        
        find_and_replace(doc, replacements)
        
        doc.save(file_name)
        messagebox.showinfo("Success", "Document saved successfully.")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# Dictionary for replacements in the Word document
replacements = {
    '[name]': tk.StringVar(value=''),
    '[standort]': tk.StringVar(value=''),
    '[id_num]': tk.StringVar(value=''),
    '[jahr]': tk.StringVar(value=''),
    '[trainer]': tk.StringVar(value=''),
    '[zeitraum_start]': tk.StringVar(value=''),
    '[zeitraum_end]': tk.StringVar(value=''),
    '[lf_num]': tk.StringVar(value=''),  # This will be set by the option menu
    '[text1]': tk.StringVar(value=''),
    '[text2]': tk.StringVar(value=''),
    '[text3]': tk.StringVar(value=''),
    '[text4]': tk.StringVar(value=''),
    '[text5]': tk.StringVar(value=''),
}

# Replace labels and entries in frame1 
info_labels = ['[name]', '[standort]', '[id_num]', '[jahr]', '[trainer]', '[zeitraum_start]', '[zeitraum_end]'] # Placeholders in word
new_labels = ['Vor -Nachname:', 'Standort:', 'Ausbildungsnachweis Nr.:', 'Ausbildungsjahr:', 'Trainer/Dozent:', 'Datum von:', 'Bis:'] # GUI

# looü and create entries
for i, label in enumerate(info_labels, start=0):
    ctk.CTkLabel(frame1, text=new_labels[i]).grid(row=i, column=0, padx=10, pady=3)
    ctk.CTkEntry(frame1, textvariable=replacements[label], width=200, height=55).grid(row=i, column=1, padx=60, pady=8)

# bitch ctkinter has issues with updating menus
""" # Function to dynamically update the option menu
def update_option_menu():
    current_options = list(notes.keys())
    lf_num_menu['menu'].delete(0, 'end')  # Clear the existing options
    for option in current_options:
        lf_num_menu['menu'].add_command(label=option, command=lambda value=option: replacements['[lf_num]'].set(value))
    if replacements['[lf_num]'].get() not in current_options:
        replacements['[lf_num]'].set(current_options[0])
    update_text_fields()


# Function to update text fields based on the selected learning field
def update_text_fields():
    selected_lf = replacements['[lf_num]'].get()
    schedule = notes.get(selected_lf, {})
    for i, day in enumerate(['montag', 'dienstag', 'mittwoch', 'donnerstag', 'freitag']):
        text_boxes[i].delete("1.0", "end")
        text_boxes[i].insert("1.0", schedule.get(day, "Keine Information")) """

# Initialize and place the option menu in the GUI
lf_num_options = list(notes.keys())  # Initial options from notes
#replacements['[lf_num]'].set(lf_num_options[0]) if lf_num_options else None
lf_num_menu = ctk.CTkOptionMenu(frame1, variable=replacements['[lf_num]'], values=lf_num_options, command=lambda _: update_text_fields())
lf_num_menu.grid(row=8, column=1, padx=60, pady=8)


#update_option_menu()

# Update Texts Button (new)
update_texts_btn = ctk.CTkButton(frame1, text="Update Texts", command=update_text_fields)
update_texts_btn.grid(row=9, column=1, padx=60, pady=8)

# Text boxes for daily notes in frame2
text_boxes = []
daily_labels = ['Montag:', 'Dienstag:', 'Mittwoch:', 'Donnerstag:', 'Freitag:']
for i, day in enumerate(daily_labels):
    ctk.CTkLabel(frame2, text=day).grid(row=i*2, column=0, sticky='nw', padx=10, pady=(10, 2))
    text_box = ctk.CTkTextbox(frame2, height=60, width=400)
    text_box.grid(row=i*2+1, column=1, padx=10, pady=(10, 10))
    text_boxes.append(text_box)

# btns ctk to save the progress
submit_btn = ctk.CTkButton(root, text="Submit", command=save_updated_document)
submit_btn.grid(row=2, column=0, columnspan=4, pady=20)

# Initialize the first tab
raise_frame(frame1)

root.mainloop()